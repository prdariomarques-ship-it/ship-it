"""GoogleDriveProvider — REST via httpx, no SDK, plus byte-level text
extraction (pypdf, python-docx). Mirrors the mocking style used for
Gmail/Calendar/Contacts.
"""
import io
from datetime import datetime, timezone
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document as DocxDocument

from providers.drive.base import (
    DriveFileTooLargeError,
    DriveProviderError,
    DriveSearchQuery,
    UnsupportedDriveFileTypeError,
)
from providers.drive.factory import UnknownDriveProviderError, get_drive_provider
from providers.drive.google.provider import (
    GoogleDriveProvider,
    _build_query,
    _extract_csv_text,
    _extract_docx_text,
    _extract_pdf_text,
    _extract_text,
    _parse_file,
    _resolve_mime_type,
)
from utils.config import get_settings


def _mock_response(json_body=None, content: bytes | None = None, status_code: int = 200) -> MagicMock:
    response = MagicMock()
    response.status_code = status_code
    response.raise_for_status = MagicMock()
    if json_body is not None:
        response.json = MagicMock(return_value=json_body)
    if content is not None:
        response.content = content
    return response


def _patch_get(get_result=None, post_result=None):
    client = MagicMock()
    if get_result is not None:
        client.get = AsyncMock(side_effect=get_result if isinstance(get_result, list) else [get_result] * 20)
    if post_result is not None:
        client.post = AsyncMock(return_value=post_result)
    client.__aenter__ = AsyncMock(return_value=client)
    client.__aexit__ = AsyncMock(return_value=False)
    return patch("providers.drive.google.provider.httpx.AsyncClient", return_value=client), client


@pytest.fixture
def provider(monkeypatch):
    monkeypatch.setattr(get_settings(), "google_client_id", "client-id")
    monkeypatch.setattr(get_settings(), "google_client_secret", "client-secret")
    monkeypatch.setattr(get_settings(), "google_drive_redirect_uri", "https://app.example.com/api/gdrive/oauth/callback")
    monkeypatch.setattr(get_settings(), "gdrive_max_file_size_bytes", 20_000_000)
    return GoogleDriveProvider()


# --- OAuth -----------------------------------------------------------------------
def test_authorization_url_includes_offline_access_readonly_scope_and_state(provider):
    url = provider.authorization_url("the-state-token")
    assert "access_type=offline" in url
    assert "prompt=consent" in url
    assert "state=the-state-token" in url
    assert "drive.readonly" in url


@pytest.mark.asyncio
async def test_exchange_code_returns_tokens(provider):
    body = {"access_token": "at1", "refresh_token": "rt1", "expires_in": 3600, "scope": "drive.readonly"}
    patcher, _ = _patch_get(post_result=_mock_response(body))
    with patcher:
        tokens = await provider.exchange_code("auth-code")
    assert tokens.access_token == "at1"
    assert tokens.refresh_token == "rt1"


@pytest.mark.asyncio
async def test_refresh_access_token_preserves_the_original_refresh_token(provider):
    body = {"access_token": "at2", "expires_in": 3600, "scope": "drive.readonly"}
    patcher, _ = _patch_get(post_result=_mock_response(body))
    with patcher:
        tokens = await provider.refresh_access_token("original-refresh-token")
    assert tokens.refresh_token == "original-refresh-token"


@pytest.mark.asyncio
async def test_token_request_failure_raises_provider_error(provider):
    import httpx as httpx_module

    client = MagicMock()
    client.post = AsyncMock(side_effect=httpx_module.ConnectError("down"))
    client.__aenter__ = AsyncMock(return_value=client)
    client.__aexit__ = AsyncMock(return_value=False)
    with patch("providers.drive.google.provider.httpx.AsyncClient", return_value=client):
        with pytest.raises(DriveProviderError):
            await provider.exchange_code("auth-code")


# --- list/search/metadata ----------------------------------------------------------
@pytest.mark.asyncio
async def test_list_files_parses_results(provider):
    body = {"files": [{"id": "f1", "name": "Contrato.pdf", "mimeType": "application/pdf", "size": "1024"}]}
    patcher, client = _patch_get(get_result=[_mock_response(body)])
    with patcher:
        files = await provider.list_files("access-token")
    assert len(files) == 1
    assert files[0].name == "Contrato.pdf"
    assert files[0].size == 1024


@pytest.mark.asyncio
async def test_search_files_sends_the_built_query(provider):
    body = {"files": []}
    patcher, client = _patch_get(get_result=[_mock_response(body)])
    with patcher:
        await provider.search_files("access-token", DriveSearchQuery(name="Contrato", folder_id="f-1"))
    sent_params = client.get.call_args.kwargs["params"]
    assert "name contains 'Contrato'" in sent_params["q"]
    assert "'f-1' in parents" in sent_params["q"]
    assert "trashed = false" in sent_params["q"]


@pytest.mark.asyncio
async def test_search_files_api_failure_raises_provider_error(provider):
    import httpx as httpx_module

    client = MagicMock()
    client.get = AsyncMock(side_effect=httpx_module.ConnectError("down"))
    client.__aenter__ = AsyncMock(return_value=client)
    client.__aexit__ = AsyncMock(return_value=False)
    with patch("providers.drive.google.provider.httpx.AsyncClient", return_value=client):
        with pytest.raises(DriveProviderError):
            await provider.search_files("access-token", DriveSearchQuery())


@pytest.mark.asyncio
async def test_get_metadata_parses_a_single_file(provider):
    body = {"id": "f1", "name": "Notas.txt", "mimeType": "text/plain", "modifiedTime": "2026-01-15T10:00:00Z"}
    patcher, _ = _patch_get(get_result=[_mock_response(body)])
    with patcher:
        metadata = await provider.get_metadata("access-token", "f1")
    assert metadata.name == "Notas.txt"
    assert metadata.modified_time == datetime(2026, 1, 15, 10, 0, tzinfo=timezone.utc)


def test_parse_file_handles_missing_size_and_modified_time():
    file = _parse_file({"id": "bare", "name": "x"})
    assert file.size == 0
    assert file.modified_time is None


# --- read_file_text: guardrails ------------------------------------------------
@pytest.mark.asyncio
async def test_read_file_text_rejects_google_native_types(provider):
    metadata_body = {"id": "f1", "name": "Planilha", "mimeType": "application/vnd.google-apps.spreadsheet"}
    patcher, _ = _patch_get(get_result=[_mock_response(metadata_body)])
    with patcher:
        with pytest.raises(UnsupportedDriveFileTypeError):
            await provider.read_file_text("access-token", "f1")


@pytest.mark.asyncio
async def test_read_file_text_rejects_unsupported_binary_types(provider):
    metadata_body = {"id": "f1", "name": "foto.png", "mimeType": "image/png"}
    patcher, _ = _patch_get(get_result=[_mock_response(metadata_body)])
    with patcher:
        with pytest.raises(UnsupportedDriveFileTypeError):
            await provider.read_file_text("access-token", "f1")


@pytest.mark.asyncio
async def test_read_file_text_rejects_oversized_files_before_downloading(provider):
    provider._max_file_size_bytes = 100  # provider already constructed; set directly (like the fixture would)
    metadata_body = {"id": "f1", "name": "grande.pdf", "mimeType": "application/pdf", "size": "999999"}
    patcher, client = _patch_get(get_result=[_mock_response(metadata_body)])
    with patcher:
        with pytest.raises(DriveFileTooLargeError):
            await provider.read_file_text("access-token", "f1")
    assert client.get.await_count == 1  # metadata only — never attempted the download


@pytest.mark.asyncio
async def test_read_file_text_downloads_and_decodes_plain_text(provider):
    metadata_body = {"id": "f1", "name": "notas.txt", "mimeType": "text/plain", "size": "10"}
    download_response = _mock_response(content="Olá, mundo!".encode())
    patcher, _ = _patch_get(get_result=[_mock_response(metadata_body), download_response])
    with patcher:
        text = await provider.read_file_text("access-token", "f1")
    assert text == "Olá, mundo!"


@pytest.mark.asyncio
async def test_read_file_text_falls_back_to_extension_for_markdown(provider):
    """Drive sometimes reports a generic mimeType for a .md file."""
    metadata_body = {"id": "f1", "name": "leia-me.md", "mimeType": "application/octet-stream", "size": "5"}
    download_response = _mock_response(content=b"# Ola")
    patcher, _ = _patch_get(get_result=[_mock_response(metadata_body), download_response])
    with patcher:
        text = await provider.read_file_text("access-token", "f1")
    assert text == "# Ola"


# --- text extraction (unit) -----------------------------------------------------
def test_extract_pdf_text_joins_pages():
    fake_reader = MagicMock()
    page1, page2 = MagicMock(), MagicMock()
    page1.extract_text.return_value = "Página 1"
    page2.extract_text.return_value = "Página 2"
    fake_reader.pages = [page1, page2]
    with patch("providers.drive.google.provider.PdfReader", return_value=fake_reader):
        text = _extract_pdf_text(b"fake-pdf-bytes")
    assert text == "Página 1\n\nPágina 2"


def test_extract_pdf_text_wraps_extraction_failures():
    with patch("providers.drive.google.provider.PdfReader", side_effect=Exception("corrupt")):
        with pytest.raises(DriveProviderError):
            _extract_pdf_text(b"not-a-real-pdf")


def test_extract_docx_text_joins_paragraphs_from_a_real_docx():
    document = DocxDocument()
    document.add_paragraph("Primeira linha")
    document.add_paragraph("Segunda linha")
    buffer = io.BytesIO()
    document.save(buffer)
    text = _extract_docx_text(buffer.getvalue())
    assert text == "Primeira linha\nSegunda linha"


def test_extract_docx_text_wraps_extraction_failures():
    with pytest.raises(DriveProviderError):
        _extract_docx_text(b"not-a-real-docx")


def test_extract_csv_text_formats_rows():
    csv_bytes = "nome,telefone\nAna,123\nBeto,456".encode()
    text = _extract_csv_text(csv_bytes)
    assert text == "nome, telefone\nAna, 123\nBeto, 456"


def test_extract_text_dispatches_by_mime_type():
    assert _extract_text("text/plain", "abc".encode()) == "abc"
    assert _extract_text("text/markdown", "# abc".encode()) == "# abc"


def test_resolve_mime_type_falls_back_to_extension_when_unrecognized():
    from providers.drive.base import DriveFile

    file = DriveFile(id="f1", name="dados.csv", mime_type="application/octet-stream")
    assert _resolve_mime_type(file) == "text/csv"


def test_resolve_mime_type_keeps_recognized_types_as_is():
    from providers.drive.base import DriveFile

    file = DriveFile(id="f1", name="qualquer.bin", mime_type="application/pdf")
    assert _resolve_mime_type(file) == "application/pdf"


def test_build_query_escapes_single_quotes():
    query = _build_query(DriveSearchQuery(name="O'Brien"))
    assert "name contains 'O\\'Brien'" in query


def test_build_query_combines_all_filters():
    query = _build_query(DriveSearchQuery(name="a", folder_id="b", mime_type="c", query="d"))
    assert "name contains 'a'" in query
    assert "'b' in parents" in query
    assert "mimeType = 'c'" in query
    assert "fullText contains 'd'" in query


# --- factory ----------------------------------------------------------------------
def test_drive_factory_resolves_google_by_default():
    get_drive_provider.cache_clear()
    assert isinstance(get_drive_provider(), GoogleDriveProvider)
    get_drive_provider.cache_clear()


def test_drive_factory_rejects_unknown_provider(monkeypatch):
    monkeypatch.setattr(get_settings(), "drive_provider", "not-a-real-provider")
    get_drive_provider.cache_clear()
    with pytest.raises(UnknownDriveProviderError):
        get_drive_provider()
    monkeypatch.setattr(get_settings(), "drive_provider", "google")
    get_drive_provider.cache_clear()
